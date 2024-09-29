import pandas as pd
from docx import Document
import os

# Carrega a lista de colaboradores do arquivo Excel
colaboradores = pd.read_excel('colaboradores.xlsx')

# Nome do arquivo de saída
nome_arquivo = 'RECIBO DE PAGAMENTO DOS DOMINGOS GRAVATAI.docx'

# Verifica se o arquivo já existe e remove se necessário
if os.path.exists(nome_arquivo):
    os.remove(nome_arquivo)

# Cria um novo documento Word
document = Document()

# Adiciona o texto do recibo para cada colaborador na mesma página
for _, row in colaboradores.iterrows():
    nome = row['nome']
    cpf = row['cpf']
    valor = row['valor']
    texto_recibo = f"""RECIBO DE PAGAMENTO
    {nome}, inscrito(a) no CPF sob o nº {cpf}, declaro para os devidos fins ter recebido nesta data, da empresa KFP SERVICE DIGITAL LTDA, inscrita no CNPJ sob o nº 41.230.154/0001-57, a importância de R${valor} concernente ao pagamento de um domingo trabalhado.

    Cachoeirinha, 29 de setembro de 2024.

    _________________________________________________
    Assinatura
    """
    document.add_paragraph(texto_recibo)
    document.add_paragraph('\n')

# Salva o arquivo de recibos
document.save(nome_arquivo)
"""
CMD:
usar a ferramenta PyInstaller para transformar seu código Python em um executável (.exe). 
>>Agora, navegue até o diretório onde seu script Python está localizado e execute o seguinte comando:
pyinstaller --onefile seu_script.py
>>Substitua seu_script.py pelo nome do seu arquivo Python. A opção --onefile faz com que todos os arquivos sejam agrupados em um único executável.
>>Depois que o PyInstaller terminar, você encontrará seu arquivo .exe na pasta dist."""