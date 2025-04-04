import pandas as pd
from docx import Document
from docx.shared import Pt
import os

# Carrega a lista de colaboradores do arquivo Excel, ignorando as primeiras 4 linhas
colaboradores = pd.read_excel('colaboradores.xlsx', skiprows=4)

# Carrega o CNPJ, razão social e data manualmente
dados_empresa = pd.read_excel('colaboradores.xlsx', nrows=3, header=None)
cnpj = dados_empresa.iloc[1, 1]  # CNPJ está na célula B2
razao_social = dados_empresa.iloc[2, 1]  # Razão social está na célula B3
data = dados_empresa.iloc[0, 2]  # Data está na célula C3
# Converte a data para o formato desejado
data = pd.to_datetime(data).strftime('%d de %B de %Y') # Exemplo: 04 de abril de 2025

# Nome do arquivo de saída
nome_arquivo = 'RECIBO DE PAGAMENTO DOS DOMINGOS GRAVATAI.docx'

# Verifica se o arquivo já existe e remove se necessário
if os.path.exists(nome_arquivo):
    os.remove(nome_arquivo)

# Cria um novo documento Word
document = Document()

# Define o tamanho da fonte para os campos
font_size = Pt(12)  # Tamanho padrão do texto
bold_font_size = Pt(14)  # Tamanho da fonte para campos em negrito

# Adiciona o texto do recibo para cada colaborador na mesma página
for _, row in colaboradores.iterrows():
    nome = row['nome']
    cpf = str(row['cpf'])  # Converte o CPF para string
    valor = str(row['valor'])  # Converte o valor para string
    
    texto = (
        f"RECIBO DE PAGAMENTO\n"
        f"{nome}, inscrito(a) no CPF sob o nº {cpf}, declaro para os devidos fins ter recebido nesta data, "
        f"da empresa {razao_social}, inscrita no CNPJ sob o nº {cnpj}, a importância de R${valor} "
        f"concernente ao pagamento de um domingo trabalhado.\n\n"
        f"Cachoeirinha, {data}.\n\n"
        f"_________________________________________________\n"
        f"Assinatura\n"
    )
    
    p = document.add_paragraph()
    run = p.add_run(texto)
    run.font.size = font_size  # Define o tamanho da fonte padrão
    document.add_paragraph('\n')

# Salva o arquivo de recibos
document.save(nome_arquivo)

"""
CMD:
usar a ferramenta PyInstaller para transformar seu código Python em um executável (.exe). 
>>Agora, navegue até o diretório onde seu script Python está localizado e execute o seguinte comando:
pyinstaller --onefile seu_script.py
>>Substitua seu_script.py pelo nome do seu arquivo Python. A opção --onefile faz com que todos os arquivos sejam agrupados em um único executável.
>>Depois que o PyInstaller terminar, você encontrará seu arquivo .exe na pasta dist."""